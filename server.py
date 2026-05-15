import os
import sys
import json
import threading
from datetime import datetime
from flask import Flask, request, jsonify, render_template
from flask_cors import CORS

sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from app.utils import load_processed_files
from app.query_rag import get_answer
from app.ingest import ingest_files, ingest_folder

app = Flask(__name__)
CORS(app)

import warnings
warnings.filterwarnings("ignore")

ingestion_status = {"status": "idle", "message": "", "files": []}

# ---------- File stats function (safe) ----------
def get_file_stats(filepath):
    """Return (page_count, word_count) for PDF/DOCX/TXT. Returns (0,0) on error."""
    pages = 0
    text = ""
    if not os.path.exists(filepath):
        return (0, 0)
    try:
        if filepath.lower().endswith('.pdf'):
            from PyPDF2 import PdfReader
            with open(filepath, 'rb') as f:
                reader = PdfReader(f)
                pages = len(reader.pages)
                for page in reader.pages:
                    text += page.extract_text() or ""
        elif filepath.lower().endswith('.docx'):
            from docx import Document
            doc = Document(filepath)
            pages = max(1, len(doc.paragraphs) // 30)
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:  # txt or other
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            pages = max(1, len(text) // 3000)
        word_count = len(text.split())
        return (pages, word_count)
    except Exception as e:
        print(f"Stats error for {filepath}: {e}")
        return (0, 0)

# ---------- Background ingestion ----------
def run_ingestion(path, is_folder, category):
    global ingestion_status
    try:
        path = path.strip().strip('"').strip("'")
        path = os.path.abspath(path)
        print(f"[INGEST] Path: {path}, exists: {os.path.exists(path)}")
        ingestion_status = {"status": "processing", "message": f"Processing {path}...", "files": []}
        if is_folder:
            ingest_folder(path, category)
        else:
            ingest_files([path], category)
        ingestion_status = {"status": "completed", "message": f"Successfully ingested: {path}", "files": []}
    except Exception as e:
        ingestion_status = {"status": "failed", "message": str(e), "files": []}
        print(f"[INGEST] Error: {e}")
        import traceback
        traceback.print_exc()

# ---------- API routes ----------
@app.route('/api/ingest', methods=['POST'])
def api_ingest():
    data = request.json
    raw_path = data.get('path')
    is_folder = data.get('is_folder', False)
    category = data.get('category', 'incoming')
    if not raw_path:
        return jsonify({'error': 'No path provided'}), 400
    thread = threading.Thread(target=run_ingestion, args=(raw_path, is_folder, category))
    thread.start()
    return jsonify({'message': 'Ingestion started in background', 'success': True})

@app.route('/api/status', methods=['GET'])
def api_status():
    return jsonify(ingestion_status)

@app.route('/api/audit', methods=['GET'])
def api_audit():
    processed = load_processed_files()
    audit_list = []
    for file_hash, original_path in processed.items():
        # Find stored copy in data/incoming/
        basename = os.path.basename(original_path)
        stored_path = None
        for root, dirs, files in os.walk('data/incoming'):
            if basename in files:
                stored_path = os.path.join(root, basename)
                break
        if not stored_path:
            stored_path = original_path
        pages, words = get_file_stats(stored_path)
        try:
            mtime = os.path.getmtime(original_path)
            date_str = datetime.fromtimestamp(mtime).isoformat()
        except:
            date_str = "unknown"
        audit_list.append({
            'hash': file_hash[:8],
            'path': original_path,
            'ingested_on': date_str,
            'pages': pages,
            'words': words
        })
    return jsonify(audit_list)

@app.route('/api/query', methods=['POST'])
def api_query():
    data = request.json
    question = data.get('question')
    if not question:
        return jsonify({'error': 'No question'}), 400
    try:
        answer = get_answer(question)
        return jsonify({'answer': answer, 'success': True})
    except Exception as e:
        return jsonify({'error': str(e), 'success': False}), 500

@app.route('/')
def index():
    return render_template('index.html')

if __name__ == '__main__':
    os.makedirs('templates', exist_ok=True)
    app.run(host='0.0.0.0', port=5000, debug=False)